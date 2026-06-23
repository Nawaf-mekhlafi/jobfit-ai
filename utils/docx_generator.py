from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import logging

class DocumentGenerator:
    @staticmethod
    def create_cover_letter_docx(cover_letter_text: str) -> BytesIO:
        """
        Generates an enterprise-formatted Word Document (.docx) strictly in-memory.
        Ensures perfect professional formatting for job applications.
        """
        logging.info("Generating in-memory DOCX file for Cover Letter...")
        
        doc = Document()
        
        # Configure enterprise document styling
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Add a professional header
        header = doc.add_heading('Cover Letter', level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()

        # Insert the AI-generated text, maintaining paragraph breaks
        paragraphs = cover_letter_text.split('\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(12) # Add clean spacing between paragraphs

        # Save to an in-memory buffer (BytesIO) instead of disk
        file_buffer = BytesIO()
        doc.save(file_buffer)
        file_buffer.seek(0) # Reset buffer pointer to the beginning
        
        logging.info("DOCX generation complete.")
        return file_buffer